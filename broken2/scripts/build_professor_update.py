"""Generate professor update Word doc summarizing project configuration.

Source: PROGRESS.md (curated), CLAUDE.md, docs/class_paper/main.tex.
Output: docs/professor_update_2026-05-08.docx
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUT = Path("/Users/chrisliang8/Desktop/safety-go2/docs/professor_update_2026-05-08.docx")


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_kv_table(doc, rows, col_widths=(2.0, 4.5)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        c0.text = ""
        c1.text = ""
        c0.paragraphs[0].add_run(k).bold = True
        c1.paragraphs[0].add_run(v)
        c0.width = Inches(col_widths[0])
        c1.width = Inches(col_widths[1])
        for cell in (c0, c1):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    return table


def add_data_table(doc, header, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for r_i, row in enumerate(rows, start=1):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    return table


def main():
    doc = Document()

    # Page setup: tighter margins for density
    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)

    # ---------- Title ----------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Adaptive Robust CBF for Quadrupedal Safety —\nProject Configuration Snapshot")
    run.bold = True
    run.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("Update for Prof. Cosner — 2026-05-07 / Target venue: CoRL 2026 (abstract May 25, paper May 28)")
    sub_run.italic = True
    sub_run.font.size = Pt(11)

    doc.add_paragraph()

    # ---------- 1. Project framing ----------
    add_heading(doc, "1. Project framing", level=1)
    add_para(
        doc,
        "Working title: Adaptive Robustness Margins for Control Barrier Functions via "
        "Privileged-Observation Reinforcement Learning — A Quadrupedal Safety-Filter Case Study.",
    )
    add_para(
        doc,
        "Core claim: a privileged-observation RL teacher can adapt the four robust-CBF parameters "
        "(α, φ, a, c) online and Pareto-dominate hand-tuned plain / ISSf / TISSf baselines on a "
        "combined fall + stuck failure metric, without sacrificing the safe-set forward-invariance "
        "structure. Inspired by Rapid Motor Adaptation (RMA): teacher conditions on privileged "
        "observations during training; a student replays it from LiDAR + base velocity + history "
        "for deploy.",
    )
    add_para(
        doc,
        "Closest published baseline: SHIELD (Yang et al. 2025, arXiv:2505.11494) — same exponential-"
        "smoothed SDF form, real LiDAR + cluster-fit-cylinder, Unitree G1 humanoid. Differentiation: "
        "(1) per-step multi-param adaptation vs their per-episode single-α calibrated by Freedman's "
        "inequality; (2) four robust slacks tied to specific uncertainty classes "
        "(Molnar / Kolathaya / Dean / boundary) plus B-fixed-X ablations as paper Table 2; "
        "(3) arbitrary obstacle shapes once perception pipeline (v2.12) lands. Cosner is a "
        "co-author on SHIELD, so this is positioned as related work, not a conflict.",
    )

    # ---------- 2. Pipeline at a glance ----------
    add_heading(doc, "2. Pipeline at a glance", level=1)
    add_kv_table(doc, [
        ("Hardware", "Unitree Go2 quadruped + Livox Mid-360 LiDAR. Go2 runs ROS 2 Humble on Jetson Orin."),
        ("Sim", "Isaac Lab on RTX 5090 (lab box, SSH chrisliang@130.64.84.163). "
                "Built on Isaac-Velocity-Flat-Unitree-Go2-v0 as locomotion base."),
        ("Locomotion checkpoint",
         "Off-the-shelf flat-trained policy (unitree_go2_flat/2026-04-15_19-24-45/exported/policy.pt). "
         "Native fall ~0.5%; in our CBF env at B0 ~20% due to 50 Hz command stream vs 10 s held training "
         "commands. Treated as a frozen, drop-in component (works with any user's locomotion stack)."),
        ("Robot model", "Single-integrator (ẋ_robot = u). L_g h ≠ 0 → standard CBF works. "
                        "HOCBF only needed if we move to torque-level control (not on roadmap)."),
        ("Outer loop", "RL teacher emits robust-CBF params (α, φ, a, c) at 50 Hz; QP wraps the "
                       "off-the-shelf locomotion controller; output goes to walking_bridge → SportClient.Move."),
        ("ROS 2 package",
         "Minimal extraction from semantic-safety: walking-bridge + LiDAR occupancy grid only. "
         "YOLO / cameras / safety filter from upstream all stripped. Source files in place; not yet "
         "built or tested on hardware."),
    ])

    # ---------- 3. CBF formulation ----------
    add_heading(doc, "3. CBF formulation and four-parameter mapping", level=1)
    add_para(
        doc,
        "Safe-set h(x) is built from per-shape analytical SDFs (boxes, cylinders, walls) combined "
        "via Eq. 19 multi-obstacle SDF + Eq. 20 exponential smoothing. Robot footprint enters as "
        "Minkowski expansion by 0.15 m. The safety QP is solved closed-form as a half-space projection "
        "on GPU — no external solver in the hot loop.",
    )
    add_para(doc, "Constraint enforced (audited 2026-05-07 against cbf_go2_env.py:168-230):", italic=True)
    add_para(doc, "    L_g h · u_safe  ≥  −α(h − c)  +  φ · ‖L_g h‖²  +  a", italic=True)
    add_para(doc, "Per-parameter mapping to the literature:", bold=True)
    add_data_table(
        doc,
        ["Param", "Range (v2.6/v2.10)", "Wide (v2.11)", "Role", "Reference"],
        [
            ["α", "[0.1, 5.0]", "unchanged", "Class-K slope on shifted h(x)−c; absorbs model / tracking error.",
             "Molnar 2021"],
            ["φ", "[0.0, 5.0]", "unchanged", "Coefficient on ‖L_g h‖² (RHS); absorbs actuation uncertainty.",
             "Kolathaya 2018 (ISSf)"],
            ["a", "[0.0, 1.0]", "[0.0, 3.0]",
             "Additive RHS slack; absorbs state-independent measurement uncertainty.",
             "Dean 2019"],
            ["c", "[0.0, 0.5]", "[0.0, 1.0]",
             "Inward shift of safe-set boundary; boundary correction when h_lidar underestimates.",
             "Boundary-correction"],
            ["b", "unused", "unused",
             "5th action slot reserved for input-dep slack b·‖u‖; needs SOCP solver — not active.",
             "Dean 2019"],
        ],
    )
    add_para(
        doc,
        "Action space is 5-D (α, φ, a, b, c) with only 4 active. Currently a known dead-weight gap: "
        "training has zero measurement noise, so a has no gradient signal and likely settles to a "
        "constant. v2.12 NoisyPerception OOD env will exercise it.",
    )
    add_para(
        doc,
        "Two L_f h items distinguished — (1) obstacle-drift term added in v2.11 behind "
        "USE_LFH_OBSTACLE_DRIFT flag (∂h/∂p_obs · v_obs added to RHS, ~15 lines, no solver change; "
        "matters most for FastObstacles); (2) HOCBF for robot model, only relevant if we switch "
        "from single- to double-integrator dynamics (kept as honest paper limitation).",
    )

    # ---------- 4. Privileged obs + network ----------
    add_heading(doc, "4. Privileged observation and teacher network", level=1)
    add_kv_table(doc, [
        ("Priv obs (PRIV-2)",
         "8207-D = 15 dynamics + 8192 occupancy grid (2 frames × 64 × 64 × 0.1 m, ego-centric, 6.4 m FOV)."),
        ("Architecture",
         "priv → CNN [Conv 2→16 s=2, 16→32 s=2, Linear → 64], dyn MLP [15 → 64] → concat → "
         "Linear 128 → 12 → Z(12) → π_teacher (128, 5)."),
        ("Bottleneck",
         "Z is 12-D — the RMA latent. get_z(obs) is exposed; student in Wk3 will reproduce Ẑ from "
         "(LiDAR, base_vel, history) via LSTM or 1-D conv, then feed frozen π_teacher."),
        ("Output", "5-D action (α, φ, a, b, c); b unused at QP."),
        ("u_des",
         "Never enters the network. Sidecar to the CBF; only used in reward via "
         "‖u_safe − u_des‖² penalty."),
    ])

    # ---------- 5. PPO + Reward ----------
    add_heading(doc, "5. PPO regularization (v2.6 working recipe — frozen)", level=1)
    add_kv_table(doc, [
        ("Optimizer", "AdamW, weight_decay = 1e-5 (monkey-patched on OnPolicyRunner.__init__)."),
        ("Entropy coef", "0.005 — load-bearing fix (5× v2.5's 0.001 prevented action-std collapse)."),
        ("Action-rate penalty", "−0.005 · ‖Δa‖² (gentler than v2.5's −0.01)."),
        ("Smoothness mechanisms",
         "Orthogonal: weight_decay smooths input → output map; action-rate smooths in time."),
        ("Episode length", "20 s; resampling controls planner switching cadence."),
    ])

    add_heading(doc, "6. Reward stack (v2.10 on disk = v2.9b reward; v2.6 ckpt was trained on v2.6 stack)",
                level=1)
    add_data_table(
        doc,
        ["Term", "Weight", "Type", "Notes"],
        [
            ["collision", "−100", "terminal", "Obstacle contact (per-shape SDF < 0)."],
            ["base_contact_penalty", "−100", "terminal", "Fall — REWARD-2 NEW; was missing before v2.9b."],
            ["stuck", "−2.0 / step", "shaping", "Triggers when ‖v_xy‖ < 0.15 m/s — REWARD-2 NEW."],
            ["infeasibility", "−10 / step", "shaping", "QP infeasibility marker."],
            ["u_safe_deviation", "−0.1 · ‖u_safe − u_des‖²", "shaping", "Per-step."],
            ["proximity", "−0.5 · exp(−min_sdf/0.5)", "shaping", "REWARD-2 halved from −1.0."],
            ["action_rate", "−0.005 · ‖Δa‖²", "shaping", "Smooths CBF params over time."],
            ["u_safe_rate", "(unregistered)", "—",
             "Function lives in code but unwired. Reserved for REWARD-3 if needed."],
        ],
    )
    add_para(doc, "v2.6 stack (paper baseline ckpt) had proximity at −1.0, no base_contact_penalty, "
                  "no stuck term. REWARD-2 (v2.9 → v2.9b → v2.10) is the retune that flipped compound "
                  "to a WIN but trades off DR width — see §10.")

    # ---------- 7. Domain randomization ----------
    add_heading(doc, "7. Domain randomization (v2.10 on disk = v2.6 narrow DR; reverted 2026-05-07)",
                level=1)
    add_data_table(
        doc,
        ["Axis", "Knob", "Train range", "OOD eval range"],
        [
            ["Friction", "static / dynamic", "(0.30, 1.20) / (0.20, 1.00)", "(0.15, 1.50) / (0.10, 1.30)"],
            ["Disturbance", "force / torque", "±10 N / ±2 Nm", "±18 N / ±3.5 Nm"],
            ["COM offset", "xy / z", "±5 cm / ±3 cm", "±8 cm / ±5 cm"],
            ["Obstacle motion", "max_speed", "0.2 m/s", "0.4 m/s"],
            ["Obstacle density", "separation_buffer", "0.4 m", "0.2 m"],
            ["Obstacle count", "K_actual ∈ [0, 20] uniform per-reset", "—", "—"],
            ["Obstacle pool", "K_MAX = 20 (8 cubes, 6 cylinders, 4 walls, 2 rect boxes, 0.20–2.0 m)",
             "—", "—"],
            ["Moving obs", "~50 % drift constant velocity per episode (±0.5 m/s per axis)", "—", "—"],
        ],
    )
    add_para(
        doc,
        "v2.7 widened DR aggressively (friction +50%, force +80%, motion 5×) at 3000 iters → "
        "under-converged. v2.8 / v2.9 / v2.9b widened mildly at 5000 iters → uniformly raised absolute "
        "combined eval values. v2.10 reverted to v2.6 narrow DR; v2.11 layers variable obstacle-motion "
        "DR (per-episode v_obs ∈ [0, 0.4] m/s) plus bimodal resample DR on top.",
    )

    # ---------- 8. Planner mix ----------
    add_heading(doc, "8. Multi-planner training mix", level=1)
    add_para(doc, "Working baseline (v2.6 ckpt — paper baseline):", bold=True)
    add_kv_table(doc, [
        ("Mix", "smooth_goal 0.40 / waypoint 0.25 / mpc 0.20 / legacy_goal 0.05 / walk 0.05 / adversarial 0.05"),
        ("Resample", "resampling_time_range = (10, 10) → 1 mid-episode switch per 20 s episode."),
    ])
    add_para(doc, "On-disk env_cfg (v2.10):", bold=True)
    add_kv_table(doc, [
        ("Mix",
         "smooth_goal 0.45 / waypoint 0.30 / mpc 0.20 / legacy_goal 0.05 — PLANNER-2b dropped walk + adversarial."),
        ("Resample", "(100, 100) → locked per episode (PLANNER-2a). Deployment-realistic."),
    ])
    add_para(doc, "v2.11 (live):", bold=True)
    add_kv_table(doc, [
        ("Resample DR",
         "Bimodal: P=0.5 → uniform [5, 15] s (mid-switch episode); P=0.5 → 100 s (locked episode). "
         "Restores v2.6's stuck-recovery regularizer that PLANNER-2a stripped in v2.8 / v2.10."),
        ("Eval", "Always locked (100 s). The training/eval mismatch is the point: training mid-switch teaches "
                 "intrinsic recovery; locked-planner eval verifies it transfers to deploy."),
    ])

    # ---------- 9. Eval matrix ----------
    add_heading(doc, "9. Headline eval matrix — v2.6 (paper baseline)", level=1)
    add_data_table(
        doc,
        ["Eval", "Type", "Margin (BR vs best baseline, combined fall+stuck)"],
        [
            ["In-distribution", "mixed", "+6.9 pp WIN"],
            ["In-dist + locked-planner eval (PLANNER-2a only)", "mixed", "+10.5 pp WIN (free win on v2.6 ckpt)"],
            ["Slippery", "priv-obs (continuous friction shift)", "+5.6 pp WIN"],
            ["DensePack", "scene-only (h(x) symmetric)", "+0.6 pp tie"],
            ["HighDisturbance", "priv-obs (episodic force/torque)", "+9.0 pp WIN"],
            ["FastObstacles", "priv-obs (grid history)", "+10.6 pp WIN"],
            ["HeavyCOM", "priv-obs (startup COM bias)", "+5.9 pp WIN"],
            ["RealisticCompound", "all 5 modest pushes simultaneously", "−0.3 pp tie"],
        ],
    )
    add_para(
        doc,
        "Pattern: WINs where the teacher has asymmetric privileged information (friction, force, "
        "COM, velocity history); ties where the playing field is symmetric (DensePack — same h(x) "
        "for both methods) or compositional (RealisticCompound — joint high-tail untrained).",
    )

    # ---------- 10. Version history ----------
    add_heading(doc, "10. Version progression (Wk2: May 4 – 10)", level=1)
    add_data_table(
        doc,
        ["Version", "Status", "Headline result", "Take-away"],
        [
            ["v2.6", "✓ paper baseline holds",
             "In-dist +6.9 pp; 5 OOD WINS / 1 compound TIE.",
             "Narrow DR + entropy 0.005 + mid-switch planner. Action-std 0.42, base_contact 5.8 % (lowest)."],
            ["v2.7", "abandoned (05-06)",
             "In-dist −0.5 pp; Slippery −11.4 pp.",
             "Wider DR + 3000 iters → under-converged checkpoint."],
            ["v2.8", "abandoned (05-06 eve)",
             "In-dist 49.3 % vs v2.6's 30.6 % (+18.7 pp regression). Stuck 7.5 % → 22.8 %.",
             "PLANNER-2a (locked train) + 2b (drop walk/adv) + mild DR widen → removed two regularizers; "
             "policy never learned to recover from CBF deflection."],
            ["v2.9", "abandoned (05-07)",
             "Stuck FIXED 22.8 → 7.7 %; fall WORSENED 26.5 → 40.8 %. −1.2 pp LOSS.",
             "REWARD-2 (base_contact_penalty −50, stuck −2.0, proximity −1.0 → −0.5). Reward-shaping "
             "direction was right; magnitude weak."],
            ["v2.9b", "abandoned (05-07 late)",
             "In-dist −0.5 pp tie; 5 OOD WINS; Compound flipped to WIN (+1.7 pp); DensePack −1.7 pp.",
             "Single-knob retune base_contact_penalty −50 → −100. Pushed fall down (40.8 → 35.0 %) but "
             "stuck rose (7.7 → 12.9 %); compound win is real."],
            ["v2.10", "✓ done (05-07 very late); fallback ckpt archived",
             "In-dist +9.9 pp WIN (combined 0.343); 4W / 2T / 1L. HeavyCOM −7.8 pp LOSS.",
             "Reverted DR + OOD ranges to v2.6; kept PLANNER-2a/2b + REWARD-2 retune. HeavyCOM "
             "mid-switch diagnostic: combined dropped 0.485 → 0.397 from eval-time planner regime alone "
             "→ smoking gun: PLANNER-2a (locked training) was the dominant HeavyCOM regression cause."],
            ["v2.11", "RUNNING (launched 05-07 overnight, ETA ~9 h)",
             "Pending.",
             "Bimodal resample DR + variable obstacle-motion DR + L_f h obstacle-drift term + "
             "WIDE_PARAM_RANGES (a [0,3], c [0,1]) + B-fixed-{α,φ,a,c} eval modes + 12 CBF training "
             "health stats + parallel 2-up eval. No new reward terms (DR-implicit shaping)."],
        ],
    )
    add_para(
        doc,
        "Decision criteria for v2.11 ship: in-dist combined ≤ 0.31 AND HeavyCOM margin ≥ +3 pp AND "
        "compound holds → v2.11 ships as paper baseline; tag + archive ckpt; proceed to Wk3 student "
        "distillation OR v2.12 perception rebuild. Fallback: v2.6 + locked-planner-eval (already "
        "+10.5 pp in-dist) is a defensible paper.",
    )

    # ---------- 11. Roadmap ----------
    add_heading(doc, "11. Remaining roadmap (CoRL 2026 — abstract May 25, paper May 28)", level=1)
    add_data_table(
        doc,
        ["Week", "Focus", "Status"],
        [
            ["Wk1 (Apr 28 – May 3)", "Multi-obstacle + varying-radius retrain + OOD eval",
             "✓ done; v1 mean rank 1.5"],
            ["Wk2 (May 4 – 10)",
             "v2 architecture + headline win",
             "v2.6 closed (5 wins / 1 tie); v2.7-v2.9b abandoned; v2.10 done, fallback. v2.11 running."],
            ["Wk3 (May 11 – 17)", "Student distillation + paper draft", "Not started; gated on ≥ v2.6 baseline."],
            ["Wk4 (May 18 – 25)", "Sim-to-real prep, polish, submit", "Not started."],
        ],
    )
    add_para(doc, "Key items remaining:", bold=True)
    add_bullet(doc, "v2.11 — bimodal resample DR + L_f h drift + wide param ranges + B-fixed-X ablations. "
                    "Tests bimodal as single attribution against v2.10's HeavyCOM regression.")
    add_bullet(doc, "v2.12 — Realistic perception in 4 layers: L1 grid-based h via distance transform "
                    "(train/deploy parity); L2 censored grid (raycast occlusion + range falloff); "
                    "L3 Gaussian noise on grid + Isaac-CBF-Go2-NoisyPerception-v0 OOD env; "
                    "L4 arbitrary obstacle shapes (irregular meshes, free once L1 lands). ~4-5 days.")
    add_bullet(doc, "Paper Table 2 — 4-axis fixed-param ablation: BR vs B-fixed-{α,φ,a,c}, each param "
                    "matched to its OOD axis (DensePack, HighDisturbance, NoisyPerception, HeavyCOM/FastObs).")
    add_bullet(doc, "Stress curve (paper figure): sweep one priv-obs axis from training-edge to ~2× past "
                    "in 4–5 steps. Shows whether BR's edge grows with difficulty.")
    add_bullet(doc, "Wk3 student distillation: re-enable LiDAR (multi-mesh K>1 obstacles), adapter "
                    "(LSTM or 1-D conv over history) → Ẑ ∈ R¹², stage-2 supervised on teacher rollouts.")
    add_bullet(doc, "Wk4 sim-to-real: TorchScript export of student adapter + π_teacher; "
                    "walking_bridge Sport-Mode → low-level Unitree motor API; 1-2 hardware demos (stretch).")

    # ---------- 12. Risks ----------
    add_heading(doc, "12. Active risks (likelihood × impact)", level=1)
    add_data_table(
        doc,
        ["Risk", "L", "I", "Mitigation / current state"],
        [
            ["Adaptation claim doesn't land", "L", "H",
             "v2.6 holds (5 wins / 1 compound tie / +6.9 pp; +10.5 pp with locked-eval). "
             "v2.9b actually won compound (+1.7 pp) where v2.6 tied — REWARD-2 fixes something on "
             "compositional stress. v2.11 testing DR-implicit retune."],
            ["Sim-to-real gap on LiDAR", "H", "M",
             "Inject matching noise during student training (v2.12 L3); "
             "v2.12 L1 closes train-deploy h-pipeline parity gap."],
            ["Reward tuning hell", "M", "M",
             "Bounded ablation budget; user preference for DR-implicit shaping over hand-crafted "
             "reward terms (cleaner paper claim)."],
            ["Teacher–student unstable", "M", "M",
             "Fall back to end-to-end PPO. Z bottleneck (12-D) keeps the surface small."],
            ["Sim-to-real gap on dynamics", "M", "M",
             "Robust CBF absorbs via φ; sim-only paper still defensible."],
            ["Locomotion fails on real floors", "M", "M",
             "Off-the-shelf flat-trained loco; gap absorbed by teacher's robust CBF."],
            ["CBF-QP infeasible at deploy", "L", "L", "Soft constraints; damping fallback."],
            ["Hardware breaks", "L", "M", "Sim-only paper still submittable."],
        ],
    )
    add_para(doc, "Drop list (cut in this order if timeline tightens):", bold=True)
    for item in [
        "Hardware demo (Wk4 stretch) → sim-only paper.",
        "L_f h obstacle-drift term + HOCBF → keep as honest future-work limitations.",
        "v2.12 NoisyPerception env → ship without it; flag a as future-work in limitations (loses 1 of 4 ablation rows).",
        "v2.11 (DR-implicit shaping + B-fixed-X) → ship v2.10 + 5-axis OOD as headline.",
        "REWARD-3 (u_safe_rate registration) → ship v2.10 with whatever fall rate it lands at.",
        "Further reward tuning after v2.10 → ship v2.10 result.",
        "v2.10 / v2.11 retrain → ship v2.6 + locked-planner-eval as the headline.",
        "Per-planner breakdown eval → headline + near-OOD covers the claim.",
    ]:
        add_bullet(doc, item)

    # ---------- 13. Discussion points ----------
    add_heading(doc, "13. Open questions for discussion", level=1)
    for q in [
        "Is bimodal resample DR (50/50 mid-switch / locked) defensible as a paper claim, or does the "
        "training/eval mismatch read as an artifact?",
        "Paper Table 2 — is per-axis fixed-param ablation (BR vs B-fixed-α/φ/a/c) the right ablation, "
        "or should we instead show param-trajectory plots from rollouts (showing the policy actually "
        "varies the params with the obs)?",
        "If v2.11 ships, do we go straight to v2.12 perception rebuild (4-5 days, hits the a-slot "
        "dead-weight gap and the deploy-realism gap simultaneously), or pivot to Wk3 student distillation?",
        "For the SHIELD comparison, do we run it as an actual baseline in our eval matrix or just "
        "position as related work?",
        "How aggressive on hardware demo for Wk4 — chair + doorway, or sim-only and full polish?",
    ]:
        add_bullet(doc, q)

    # ---------- Footer note ----------
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot_run = foot.add_run(
        "Generated 2026-05-07. Source: PROGRESS.md (1191 lines), CLAUDE.md, docs/class_paper/main.tex. "
        "v2.11 results pending (~9 h ETA)."
    )
    foot_run.italic = True
    foot_run.font.size = Pt(9)
    foot_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
